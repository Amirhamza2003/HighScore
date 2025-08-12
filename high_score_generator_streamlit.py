import streamlit as st
import os
import time
import requests
from docx import Document
from groq import Groq
import pandas as pd
import base64
from datetime import datetime
import json

# Page configuration
st.set_page_config(
    page_title="High Score Question Generator",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for professional styling
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 10px;
        margin-bottom: 2rem;
        text-align: center;
        color: white;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    
    .section-header {
        background: linear-gradient(90deg, #f093fb 0%, #f5576c 100%);
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        color: white;
        font-weight: bold;
    }
    
    .info-box {
        background: linear-gradient(90deg, #a8edea 0%, #fed6e3 100%);
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        border-left: 4px solid #667eea;
    }
    
    .success-box {
        background: linear-gradient(90deg, #d4fc79 0%, #96e6a1 100%);
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        border-left: 4px solid #28a745;
    }
    
    .warning-box {
        background: linear-gradient(90deg, #ffecd2 0%, #fcb69f 100%);
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        border-left: 4px solid #ffc107;
    }
    
    .stButton > button {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        padding: 0.75rem 2rem;
        border-radius: 25px;
        font-weight: bold;
        font-size: 1.1rem;
        transition: all 0.3s ease;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 12px rgba(0, 0, 0, 0.15);
    }
    
    .curriculum-example {
        background: #f8f9fa;
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #dee2e6;
        font-family: 'Courier New', monospace;
        font-size: 0.9rem;
        margin: 0.5rem 0;
    }
    
    .progress-container {
        background: #f8f9fa;
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #dee2e6;
        margin: 1rem 0;
    }
    
    .question-output {
        background: #ffffff;
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #dee2e6;
        margin: 0.5rem 0;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.05);
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'questions_generated' not in st.session_state:
    st.session_state.questions_generated = []
if 'generation_in_progress' not in st.session_state:
    st.session_state.generation_in_progress = False

def main():
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>🎯 HIGH SCORE QUESTION GENERATOR</h1>
        <p style="font-size: 1.2rem; margin: 0;">AI-Powered Math Assessment Question Creation</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar for configuration
    with st.sidebar:
        st.markdown("""
        <div class="section-header">
            ⚙️ CONFIGURATION
        </div>
        """, unsafe_allow_html=True)
        
        # API Key
        api_key = st.text_input(
            "🔑 Groq AI API Key",
            type="password",
            help="Enter your Groq AI API key"
        )
        
        # Difficulty Level
        difficulty = st.selectbox(
            "🎯 Difficulty Level",
            ["easy", "moderate", "hard", "mix"],
            index=1,
            help="Select the difficulty level for generated questions"
        )
        
        # Input Method
        input_method = st.radio(
            "📚 Input Method",
            ["Manual Input", "Excel Import"],
            help="Choose how to input curriculum data"
        )
        
        st.markdown("---")
        
        # Quick Start Guide
        st.markdown("""
        <div class="info-box">
            <h4>🚀 Quick Start Guide</h4>
            <ol>
                <li>Enter your base questions</li>
                <li>Select difficulty level</li>
                <li>Choose input method</li>
                <li>Enter curriculum data</li>
                <li>Click Generate Questions</li>
            </ol>
        </div>
        """, unsafe_allow_html=True)
    
    # Main content area
    col1, col2 = st.columns([2, 1])
    
    with col1:
        # Base Questions Section
        st.markdown("""
        <div class="section-header">
            📝 BASE QUESTIONS
        </div>
        """, unsafe_allow_html=True)
        
        base_q1 = st.text_area(
            "Base Question 1",
            height=120,
            placeholder="Enter your first base question here...",
            help="This will be used as a reference for generating similar questions"
        )
        
        base_q2 = st.text_area(
            "Base Question 2", 
            height=120,
            placeholder="Enter your second base question here...",
            help="This will be used as a reference for generating similar questions"
        )
        
        # Curriculum Section
        st.markdown("""
        <div class="section-header">
            📚 CURRICULUM CONFIGURATION
        </div>
        """, unsafe_allow_html=True)
        
        if input_method == "Manual Input":
            st.markdown("""
            <div class="info-box">
                <h4>📋 Manual Input Format</h4>
                <p>Enter curriculum paths in the format: <strong>Subject | Unit | Topic</strong></p>
                <p>Separate multiple paths with semicolons (;)</p>
            </div>
            """, unsafe_allow_html=True)
            
            # Example curriculum data
            example_curriculum = """Quantitative Math | Problem Solving | Numbers and Operations; Quantitative Math | Algebra | Quadratic Equations & Functions; Quantitative Math | Geometry and Measurement | Area & Volume"""
            
            st.markdown("""
            <div class="curriculum-example">
                <strong>Example:</strong><br>
                """ + example_curriculum + """
            </div>
            """, unsafe_allow_html=True)
            
            curriculum_input = st.text_area(
                "Curriculum Paths",
                value=example_curriculum,
                height=100,
                help="Enter curriculum paths in the format: Subject | Unit | Topic"
            )
            
        else:
            st.markdown("""
            <div class="info-box">
                <h4>📊 Excel Import</h4>
                <p>Upload an Excel file with 3 columns: Subject, Unit, Topic</p>
            </div>
            """, unsafe_allow_html=True)
            
            uploaded_file = st.file_uploader(
                "Choose Excel file",
                type=['xlsx', 'xls'],
                help="Upload Excel file with Subject, Unit, Topic columns"
            )
            
            if uploaded_file is not None:
                try:
                    df = pd.read_excel(uploaded_file)
                    if len(df.columns) >= 3:
                        st.success(f"✅ File uploaded successfully! Found {len(df)} curriculum paths.")
                        st.dataframe(df.head(), use_container_width=True)
                        curriculum_input = df
                    else:
                        st.error("❌ Excel file must have at least 3 columns: Subject, Unit, Topic")
                        curriculum_input = None
                except Exception as e:
                    st.error(f"❌ Error reading Excel file: {e}")
                    curriculum_input = None
            else:
                curriculum_input = None
        
        # Generate Button
        st.markdown("<br>", unsafe_allow_html=True)
        
        if st.button("🚀 GENERATE QUESTIONS", type="primary", use_container_width=True):
            if not base_q1 or not base_q2:
                st.error("❌ Please enter both base questions")
            elif curriculum_input is None or (isinstance(curriculum_input, str) and not curriculum_input.strip()) or (isinstance(curriculum_input, pd.DataFrame) and curriculum_input.empty):
                st.error("❌ Please provide curriculum data")
            elif not api_key:
                st.error("❌ Please enter your Groq API key")
            else:
                st.session_state.generation_in_progress = True
                generate_questions(base_q1, base_q2, curriculum_input, difficulty, api_key)
    
    with col2:
        # Progress and Output Section
        st.markdown("""
        <div class="section-header">
            📊 PROGRESS & OUTPUT
        </div>
        """, unsafe_allow_html=True)
        
        if st.session_state.generation_in_progress:
            with st.spinner("🔄 Generating questions..."):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                # Simulate progress
                for i in range(100):
                    time.sleep(0.01)
                    progress_bar.progress(i + 1)
                    if i < 30:
                        status_text.text("🔄 Initializing...")
                    elif i < 60:
                        status_text.text("🤖 Generating questions with AI...")
                    elif i < 90:
                        status_text.text("📝 Creating Word document...")
                    else:
                        status_text.text("✅ Complete!")
                
                progress_bar.empty()
                status_text.empty()
                st.success("🎉 Questions generated successfully!")
        
        # Display generated questions
        if st.session_state.questions_generated:
            st.markdown("""
            <div class="section-header">
                📋 GENERATED QUESTIONS
            </div>
            """, unsafe_allow_html=True)
            
            for i, question_data in enumerate(st.session_state.questions_generated):
                with st.expander(f"Topic {i+1}: {question_data['topic']}", expanded=True):
                    st.markdown(f"""
                    <div class="question-output">
                        <strong>Subject:</strong> {question_data['subject']}<br>
                        <strong>Unit:</strong> {question_data['unit']}<br>
                        <strong>Topic:</strong> {question_data['topic']}<br>
                        <strong>Difficulty:</strong> {question_data['difficulty']}<br>
                        <br>
                        {question_data['content']}
                    </div>
                    """, unsafe_allow_html=True)
            
            # Download button
            if st.button("📥 Download Word Document", use_container_width=True):
                download_word_document(st.session_state.questions_generated)

def parse_curriculum_input(curriculum_input):
    """Parse curriculum input from manual text or Excel data"""
    if isinstance(curriculum_input, str):
        # Manual input
        paths = []
        entries = [e.strip() for e in curriculum_input.split(';') if e.strip()]
        for entry in entries:
            parts = [p.strip() for p in entry.split('|')]
            if len(parts) == 3:
                paths.append({
                    'subject': parts[0],
                    'unit': parts[1],
                    'topic': parts[2]
                })
        return paths
    elif isinstance(curriculum_input, pd.DataFrame):
        # Excel input
        paths = []
        for _, row in curriculum_input.iterrows():
            paths.append({
                'subject': str(row.iloc[0]).strip(),
                'unit': str(row.iloc[1]).strip(),
                'topic': str(row.iloc[2]).strip()
            })
        return paths
    return []

def generate_questions(base_q1, base_q2, curriculum_input, difficulty, api_key):
    """Generate questions using Groq API"""
    try:
        # Parse curriculum paths
        curriculum_paths = parse_curriculum_input(curriculum_input)
        
        if not curriculum_paths:
            st.error("❌ No valid curriculum paths found")
            return
        
        # Set API key
        os.environ["GROQ_API_KEY"] = api_key
        groq_client = Groq(api_key=api_key)
        
        st.session_state.questions_generated = []
        
        # Generate questions for each topic
        for path in curriculum_paths:
            with st.spinner(f"Generating questions for: {path['topic']}"):
                questions = generate_questions_for_topic(
                    groq_client, base_q1, base_q2, path, difficulty
                )
                
                if questions:
                    st.session_state.questions_generated.append({
                        'subject': path['subject'],
                        'unit': path['unit'],
                        'topic': path['topic'],
                        'difficulty': difficulty,
                        'content': questions
                    })
                    st.success(f"✅ Generated 2 questions for: {path['topic']}")
                else:
                    st.error(f"❌ Failed to generate questions for: {path['topic']}")
        
        st.session_state.generation_in_progress = False
        
    except Exception as e:
        st.error(f"❌ Error generating questions: {e}")
        st.session_state.generation_in_progress = False

def generate_questions_for_topic(groq_client, base_q1, base_q2, path, difficulty):
    """Generate questions for a specific topic"""
    # Determine difficulty for this topic
    if difficulty == 'mix':
        import random
        topic_difficulty = random.choice(['easy', 'moderate', 'hard'])
    else:
        topic_difficulty = difficulty
    
    prompt = f'''
You are a math education expert. Given the following base questions:

1. {base_q1}
2. {base_q2}

Create 2 new math questions similar in style and difficulty, using the following format:

title: Assessment title for {path['topic']}
description: assessment description for {path['topic']}
question: Write your question here
instruction: Write instruction here
difficulty: {topic_difficulty}
order: Question number
option A: ...
option B: ...
option C: ...
option D: ...
option E: ...
correct_answer: A/B/C/D/E
explanation: Write your question explanation here
subject: {path['subject']}
unit: {path['unit']}
topic: {path['topic']}
plusmarks: 1

Make sure the questions are similar in structure and difficulty to the base questions, and are relevant to the topic: {path['topic']}.
Generate exactly 2 questions for this topic.
'''
    
    try:
        response = groq_client.chat.completions.create(
            model="llama3-8b-8192",
            messages=[
                {"role": "system", "content": "You are a math education expert who creates high-quality assessment questions."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Error generating questions with Groq: {e}")
        return None

def download_word_document(questions_data):
    """Create and download Word document"""
    try:
        doc = Document()
        doc.add_heading('High Score Question Generator - Math Assessment', 0)
        doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('')
        
        for i, question_data in enumerate(questions_data, 1):
            doc.add_heading(f'Topic {i}: {question_data["topic"]}', level=1)
            doc.add_paragraph(f'Subject: {question_data["subject"]}')
            doc.add_paragraph(f'Unit: {question_data["unit"]}')
            doc.add_paragraph(f'Difficulty: {question_data["difficulty"]}')
            doc.add_paragraph('')
            doc.add_paragraph(question_data['content'])
            doc.add_paragraph('')
        
        # Save document
        filename = f'high_score_questions_{int(time.time())}.docx'
        doc.save(filename)
        
        # Read file and create download button
        with open(filename, "rb") as file:
            btn = st.download_button(
                label="📥 Click to Download",
                data=file.read(),
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        # Clean up
        os.remove(filename)
        
    except Exception as e:
        st.error(f"❌ Error creating Word document: {e}")

if __name__ == "__main__":
    main()
